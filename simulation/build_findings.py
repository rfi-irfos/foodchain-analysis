#!/usr/bin/env python3
"""
Build: Simulation Findings Report — Ana Diez ABM
RFI-IRFOS Research · 2026-06-14
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

CEL   = RGBColor(0x4A, 0x90, 0xD9)   # celeste
NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
GOLD  = RGBColor(0xF5, 0xA6, 0x23)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x5A, 0x7A, 0x9A)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def para(text="", bold=False, italic=False, size=11, color=None, align=None,
         space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color or NAVY
    return p

def heading(text, level=1):
    if level == 1:
        p = para(text, bold=True, size=15, color=NAVY, space_before=14, space_after=4)
        # underline via border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '4A90D9')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        para(text, bold=True, size=12, color=CEL, space_before=10, space_after=2)
    elif level == 3:
        para(text, bold=True, size=11, color=GREY, space_before=6, space_after=2)

def bullet(label, value, note=None, value_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.6)
    r1 = p.add_run("  •  ")
    r1.font.color.rgb = CEL
    r1.font.size = Pt(11)
    r2 = p.add_run(label + "  ")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY
    r3 = p.add_run(value)
    r3.font.size = Pt(11)
    r3.font.color.rgb = value_color or GREEN
    if note:
        r4 = p.add_run("  " + note)
        r4.italic = True
        r4.font.size = Pt(10)
        r4.font.color.rgb = GREY

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8E0F4')
    pBdr.append(bottom)
    pPr.append(pBdr)

def callout(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '4A90D9' if not color else color)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREY

# ═══════════════════════════════════════════════════════════════════════════
#  TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run("Simulation Findings Report")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("Neurobiological-Fitness Consequence Separation ABM")
r2.font.size = Pt(12)
r2.italic = True
r2.font.color.rgb = CEL

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(12)
r3 = p3.add_run("Ana Diez (2026)  ·  Independent replication + exploration  ·  RFI-IRFOS Research  ·  2026-06-14")
r3.font.size = Pt(9)
r3.font.color.rgb = GREY

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
heading("Executive Summary", 1)
para(
    "We independently re-implemented Ana Diez's ABM (100 agents, 8 rules, synchronous update) "
    "from the paper specification and ran a full battery of experiments: baseline simulation, "
    "null-model dissection, Latin Hypercube Sampling (120 samples), Sobol variance decomposition "
    "(448 evaluations), and Approximate Bayesian Computation (200 simulations). "
    "The core thesis holds robustly. Several secondary findings sharpen, complicate, or extend "
    "the original paper's claims.",
    size=11, color=NAVY, space_after=4
)
callout(
    "Bottom line: humanity is indeed on a collapsing trajectory in this model — and no single "
    "parameter intervention reverses it. The model has no stable attractor. It only falls, "
    "at different speeds."
)

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 1 — BASELINE
# ═══════════════════════════════════════════════════════════════════════════
heading("1. Baseline Simulation (300 ticks, default parameters)", 1)
para("Default parameters as specified in the paper. 1 tick = 1 year, starting 1970.", size=10, color=GREY, space_after=4)

bullet("Earth Overshoot Day (final, year 2270):", "72.6  (March 13th)",
       note="started at 365 in 1970")
bullet("Consequence separation (final):", "4.97",
       note="0.000 at tick 0 → monotonic growth, never reverses")
bullet("Reproductive physiology (final):", "0.0011",
       note="near-total physiological collapse, not a gentle decline")
bullet("Synaptic habituation (final mean):", "35.4×",
       note="vs 1.0 at initialisation")
bullet("ABC RMSE (1970–2023, 54 ticks):", "22.95 days",
       note="Ana's paper: 5.64 days — gap discussed in Section 5", value_color=GOLD)
bullet("Reproductive decline rate:", "−2.25%/year",  note="compound", value_color=RED)

para(
    "The EOD trajectory is monotonically decreasing throughout the 300-year run with no "
    "sign of stabilisation. Consequence separation grows at an accelerating rate driven by "
    "synaptic habituation compounding annually.",
    size=10.5, color=NAVY, space_before=6, space_after=4
)

divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 2 — ROBUSTNESS (LHS)
# ═══════════════════════════════════════════════════════════════════════════
heading("2. Robustness — Latin Hypercube Sampling (120 samples, 50 ticks)", 1)
para("Parameters randomly sampled across their full plausible ranges via LHS.", size=10, color=GREY, space_after=4)

bullet("Consequence separation positive:", "100%",  note="Ana's paper: 100%")
bullet("Ecological overshoot (EOD < Dec 31):", "100%", note="Ana's paper: 100%")
bullet("Reproductive physiology decline:", "100%",  note="Ana's paper: 97.6%  — we exceed her finding")
bullet("CS range across samples:", "0.029 – 3.309",  note="mean 1.470")

callout(
    "The model is parameter-robust. No combination of inputs tested across the LHS space "
    "avoids overshoot or a positive consequence separation. The thesis holds globally, "
    "not just at calibrated defaults."
)
divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 3 — SOBOL
# ═══════════════════════════════════════════════════════════════════════════
heading("3. Sobol Variance Decomposition — What Actually Drives the Model?", 1)
para("Saltelli 2002 estimator, 448 evaluations. Output variable: consequence_separation at final tick.", size=10, color=GREY, space_after=4)

heading("Ranked by first-order index (S1):", 3)
bullet("resource_acquisition_drive:", "S1 = 0.850  ·  ST = 0.853",
       note="dominates — 85% of all variance")
bullet("base_habituation_rate:", "S1 = 0.291  ·  ST = 0.347",
       note="Ana's central parameter — second place")
bullet("status_consumption_conversion:", "S1 = 0.003  ·  ST = 0.000",
       note="negligible")
bullet("base_status_impulse_sensitivity:", "S1 = 0.000  ·  ST = 0.000",
       note="effectively zero")
bullet("cultural_selection_rate:", "S1 = 0.000  ·  ST = 0.000",
       note="effectively zero")

callout(
    "The paper frames synaptic habituation (base_habituation_rate) as the central causal "
    "mechanism. The Sobol analysis challenges this: resource_acquisition_drive explains 3× "
    "more variance. Habituation matters — but what you are born wanting matters more. "
    "This suggests the paper's narrative emphasis may deserve revision."
)
divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 4 — PATHWAY DISSECTION
# ═══════════════════════════════════════════════════════════════════════════
heading("4. Reproductive Physiology — HPA vs. EDC Pathway Dissection", 1)
para("150-tick runs, ablating each biological pathway independently.", size=10, color=GREY, space_after=4)

bullet("Baseline (both pathways active):", "repro = 0.133")
bullet("HPA knocked out (stress axis only):", "repro = 0.273",  note="Δ +0.140  (+105%)", value_color=GOLD)
bullet("EDC knocked out (chemical pathway only):", "repro = 0.493", note="Δ +0.360  (+270%)", value_color=GOLD)
bullet("Both knocked out:", "repro = 1.000", note="perfect physiological maintenance")

callout(
    "EDC accumulation is the dominant reproductive suppressor — contributing roughly 2.6× "
    "more damage than the HPA-HPG stress pathway. The paper treats both pathways as roughly "
    "equal contributors. The numbers disagree. This may be the most actionable finding "
    "for future model refinement: EDC delay time and coefficient deserve a dedicated "
    "sensitivity analysis of their own."
)
divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 5 — NULL MODEL
# ═══════════════════════════════════════════════════════════════════════════
heading("5. Null Model — Is the Mechanism Necessary or Sufficient?", 1)
para("hab_rate set to 1.000 (no growth, no decoupling). 300 ticks.", size=10, color=GREY, space_after=4)

bullet("CS (with mechanism, hab=1.012):", "4.97")
bullet("CS (null model, hab=1.000):", "2.81", note="−43%", value_color=GOLD)
bullet("EOD final — both models:", "72.6", note="identical floors")
bullet("Repro final — both models:", "0.0011", note="identical collapse")

callout(
    "Consequence separation exists even without habituation — at 57% of the full-mechanism "
    "level. This means the mechanism is not strictly necessary to produce separation: it is "
    "sufficient, and it amplifies. This is actually a stronger result for the thesis than "
    "the paper implies: the claim should be reframed from 'the mechanism produces overshoot' "
    "to 'the mechanism dramatically amplifies an already-existing separation.' "
    "The null model finding supports, not undermines, the core argument."
)
divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 6 — INTERVENTIONS
# ═══════════════════════════════════════════════════════════════════════════
heading("6. Intervention Experiments", 1)

heading("6a. Information transparency (info_suppression sweep)", 2)
bullet("Full suppression (1.50):", "repro = 0.064  ·  EOD = 138.2")
bullet("Full transparency (0.00):", "repro = 0.273  ·  EOD = 138.2",
       note="4.3× repro improvement, zero EOD effect")
callout(
    "Information transparency saves reproductive physiology but does not move EOD by a "
    "single day. People can know everything about ecological collapse and still consume the "
    "same. This is arguably the paper's most striking emergent finding — and it is not "
    "explicitly highlighted in the original text."
)

heading("6b. Cultural transmission — wellbeing prestige (wpw sweep)", 2)
para("wellbeing_prestige_weight swept 0.0 → 1.0 in steps of 0.25. 150 ticks.", size=10, color=GREY, space_after=4)
bullet("Effect on EOD:", "zero", note="all values produce identical EOD = 138.2", value_color=RED)
bullet("Effect on CS:", "zero", note="all values identical", value_color=RED)
bullet("Effect on repro:", "zero", note="all values identical", value_color=RED)
callout(
    "Rule 8 (bidirectional cultural transmission) is effectively inert in this model. "
    "Whether the population copies high-consumption elites or high-wellbeing-efficiency "
    "agents makes no difference to any outcome variable. The resource_acquisition_drive "
    "overwhelms cultural rebalancing before it can propagate. "
    "This is either a fundamental insight (cultural nudges are too weak against biological "
    "drives) or a model design issue worth revisiting."
)

heading("6c. Best-case scenario", 2)
para("Minimum rad (0.002), maximum wellbeing prestige (1.0), full transparency (0.0), minimum status sensitivity (0.05). 300 ticks.", size=10, color=GREY, space_after=4)
bullet("EOD final:", "215.2", note="vs 72.6 baseline — meaningfully slower collapse")
bullet("EOD drift (last 50 ticks):", "−17.9 days", note="still declining, no stabilisation", value_color=RED)
bullet("Consequence separation:", "1.653", note="vs 4.97 baseline")
bullet("Repro final:", "0.136", note="vs 0.001 baseline — major improvement", value_color=GOLD)
callout(
    "The best-case scenario slows the trajectory significantly but finds no stable equilibrium. "
    "The model has no attractor state other than collapse at variable speed. "
    "There is no recoverable trajectory in the current design."
)
divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 7 — ABC
# ═══════════════════════════════════════════════════════════════════════════
heading("7. ABC Parameter Estimation", 1)
para("200 simulations, 5th percentile rejection, fitted against EOD 1970–2023.", size=10, color=GREY, space_after=4)

bullet("Best-fit RMSE:", "18.17 days", note="vs Ana's paper: 5.64 days", value_color=GOLD)
bullet("Accepted simulations:", "10 / 200", note="5th percentile = 10")
bullet("Posterior: base_habituation_rate:", "mean 1.0198  ·  95% CI [1.007, 1.029]",
       note="Ana's paper: ~1.012")
bullet("Posterior: status_impulse_sensitivity:", "mean 0.856  ·  95% CI [0.764, 0.983]",
       note="vs default 0.30 — ABC pulls 3× higher")
bullet("Posterior: status_consumption_conversion:", "mean 1.814  ·  95% CI [1.534, 1.984]",
       note="vs default 0.50 — ABC pulls 3.6× higher")
bullet("Simulated EOD 2023:", "245.0", note="empirical: 209 — 36-day gap")

callout(
    "The RMSE gap (18 vs 5.64) has two likely causes. First: we ran 10× fewer ABC simulations "
    "(200 vs 2000), giving coarser posterior coverage. Second: the commons_health formula and "
    "EOD calculation were reconstructed from specification since the original Python was not "
    "available — small formula differences compound over 54 years. "
    "Crucially, the ABC posteriors reveal that the model fits historical data best when "
    "status pressure is far stronger than the paper's defaults assume. This is a legitimate "
    "research question: is the default parameterisation underestimating social contagion effects?"
)
divider()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 8 — SUMMARY TABLE
# ═══════════════════════════════════════════════════════════════════════════
heading("8. Summary of Findings", 1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Finding"
hdr[1].text = "Result"
hdr[2].text = "Implication"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A3A5C')
    tcPr.append(shd)

rows = [
    ("Core thesis robustness (LHS)", "100% overshoot, 100% CS positive across all 120 samples", "Thesis holds globally"),
    ("Dominant parameter (Sobol)", "resource_acquisition_drive S1=0.85", "Narrative should lead with this, not habituation"),
    ("Habituation role (Sobol)", "S1=0.29, second place", "Important but not dominant"),
    ("EDC vs HPA", "EDC = 2.6× more damage than HPA", "EDC pathway deserves dedicated analysis"),
    ("Null model", "CS = 2.81 without mechanism (vs 4.97 with)", "Reframe: mechanism amplifies, not generates"),
    ("Cultural transmission (Rule 8)", "Zero effect across all wellbeing_prestige_weight values", "Either a key insight or a model design issue"),
    ("Information transparency", "4.3× repro gain, zero EOD gain", "Knowledge does not reduce consumption"),
    ("Stable equilibrium", "None found — model always falling", "No recoverable trajectory in current design"),
    ("ABC RMSE", "18.2 days vs Ana's 5.64", "Reconstruction gap — needs original Python to close"),
    ("Repro collapse", "Final value 0.0011 at 300 ticks", "Framing as 'collapse', not 'decline', is warranted"),
]

for finding, result, impl in rows:
    row = table.add_row().cells
    row[0].text = finding
    row[1].text = result
    row[2].text = impl
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = NAVY

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
#  FOOTER NOTE
# ═══════════════════════════════════════════════════════════════════════════
divider()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run(
    "Generated by RFI-IRFOS Research · rfi.irfos@gmail.com · 2026-06-14  ·  "
    "Full simulation source: sim.py + dashboard.html (python3 sim.py → http://localhost:8765)"
)
r.font.size = Pt(8.5)
r.italic = True
r.font.color.rgb = GREY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = "/home/eri-irfos/Desktop/Anas Whitepaper/Simulation Findings — Ana Diez ABM.docx"
doc.save(out)
print(f"saved: {out}")
